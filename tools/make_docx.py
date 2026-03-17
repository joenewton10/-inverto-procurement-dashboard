#!/usr/bin/env python3
"""
Generate styled DOCX files from application pipeline .md outputs.
Style: Georgia, steel-blue (#1F497D) headings, open-circle bullets, two-column skills table.
Skips KEYWORD GAP ANALYSIS block automatically.

Usage: /c/Python313/python.exe tools/make_docx.py <output_folder_path>
Produces: 02-resume.docx  03-cover-letter.docx
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Constants ─────────────────────────────────────────────
DIVIDER_CHAR = "\u2500"          # ─
EM_DASH      = "\u2014"          # —
OPEN_CIRCLE  = "\u25CB"          # ○
BLUE         = RGBColor(0x1F, 0x49, 0x7D)
BLACK        = RGBColor(0x00, 0x00, 0x00)
FONT         = "Georgia"
SKIP_SECTIONS = {"KEYWORD GAP ANALYSIS"}

# ── Helpers ───────────────────────────────────────────────

def set_margins(doc, inches=1.25):
    for section in doc.sections:
        section.top_margin    = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin   = Inches(inches)
        section.right_margin  = Inches(inches)


def set_default_font(doc, size=10.5):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(size)


def sp(p, before=0, after=4):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)


def add_run(p, text, size=10.5, bold=False, italic=False, underline=False, color=BLACK):
    r = p.add_run(text)
    r.font.name     = FONT
    r.font.size     = Pt(size)
    r.bold          = bold
    r.italic        = italic
    r.underline     = underline
    r.font.color.rgb = color
    return r


def is_divider(line):
    s = line.strip()
    return bool(s) and all(c == DIVIDER_CHAR for c in s)


# ── Resume block builders ─────────────────────────────────

def para_name(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, size=22, bold=False, color=BLUE)
    sp(p, 0, 6)


def para_contact(doc, text):
    """Pipe-separated contact line — render centered, plain."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, size=11, color=BLACK)
    sp(p, 0, 3)


def para_section_heading(doc, text):
    p = doc.add_paragraph()
    add_run(p, text.upper(), size=11, bold=True, underline=True, color=BLUE)
    sp(p, 12, 5)


def para_job_title(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size=11, bold=True, color=BLACK)
    sp(p, 8, 2)


def para_subtitle(doc, text):
    """Thesis subtitle or sub-note — italic, smaller."""
    p = doc.add_paragraph()
    add_run(p, text, size=10, italic=True, color=BLACK)
    sp(p, 0, 3)


def para_date_location(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size=10.5, italic=True, color=BLACK)
    sp(p, 0, 3)


def para_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Inches(0.3)
    p.paragraph_format.first_line_indent  = Inches(-0.2)
    add_run(p, OPEN_CIRCLE + "  ", size=10.5, color=BLACK)
    add_run(p, text, size=10.5, color=BLACK)
    sp(p, 0, 3)


def para_bullet_with_bold_label(doc, text):
    """Bullet where text may start with **Bold label** — rest."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    add_run(p, OPEN_CIRCLE + "  ", size=10.5, color=BLACK)
    m = re.match(r"\*\*(.+?)\*\*\s*[" + EM_DASH + r"\-]+\s*(.*)", text)
    if m:
        add_run(p, m.group(1), size=10.5, bold=True, color=BLACK)
        add_run(p, " " + EM_DASH + " " + m.group(2), size=10.5, color=BLACK)
    else:
        add_run(p, text, size=10.5, color=BLACK)
    sp(p, 3, 3)


def table_two_col_skills(doc, skills):
    """Render skills list as a borderless two-column table."""
    half       = (len(skills) + 1) // 2
    left_items = skills[:half]
    right_items = skills[half:]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove all borders
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    tblBdr = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBdr.append(b)
    tblPr.append(tblBdr)

    def fill_cell(cell, items):
        # Remove the default empty paragraph word inserts
        for existing_p in cell.paragraphs:
            existing_p._element.getparent().remove(existing_p._element)
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            add_run(p, OPEN_CIRCLE + "  ", size=10.5, color=BLACK)
            add_run(p, item, size=10.5, color=BLACK)
            sp(p, 0, 3)

    row = table.rows[0]
    fill_cell(row.cells[0], left_items)
    fill_cell(row.cells[1], right_items)
    sp(doc.add_paragraph(), 0, 4)   # spacer after table


# ── Resume generator ──────────────────────────────────────

def generate_resume_docx(md_path, docx_path):
    raw_lines = Path(md_path).read_text(encoding="utf-8").splitlines()

    # Skip everything up to and including the "---" separator
    # (removes KEYWORD GAP ANALYSIS block)
    start_idx = 0
    for idx, line in enumerate(raw_lines):
        if line.strip() == "---":
            start_idx = idx + 1
            break
    lines = raw_lines[start_idx:]

    doc = Document()
    set_margins(doc)
    set_default_font(doc)

    i = 0

    # ── Header (name + contact, before first divider) ──────
    header_lines = []
    while i < len(lines):
        stripped = lines[i].strip()
        if is_divider(lines[i]):
            i += 1
            break
        if stripped:
            header_lines.append(stripped)
        i += 1

    if header_lines:
        para_name(doc, header_lines[0])
        for cl in header_lines[1:]:
            para_contact(doc, cl)

    # ── Body ───────────────────────────────────────────────
    current_section  = None
    next_is_heading  = True
    skip_section     = False
    skills_buffer    = []   # accumulate KEY SKILLS bullets for two-col table
    last_was_date    = False  # track position in PROFESSIONAL EXPERIENCE

    def flush_skills():
        nonlocal skills_buffer
        if skills_buffer:
            table_two_col_skills(doc, skills_buffer)
            skills_buffer = []

    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()

        # Divider
        if is_divider(line):
            if current_section == "KEY SKILLS":
                flush_skills()
            next_is_heading = True
            last_was_date   = False
            i += 1
            continue

        # Ignore stray "---"
        if stripped == "---":
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # ── Section heading ────────────────────────────────
        if next_is_heading:
            sec = stripped.upper()
            if current_section == "KEY SKILLS":
                flush_skills()
            skip_section    = sec in SKIP_SECTIONS
            current_section = sec
            if not skip_section:
                para_section_heading(doc, stripped)
            next_is_heading = False
            last_was_date   = False
            i += 1
            continue

        if skip_section:
            i += 1
            continue

        # ── KEY SKILLS — collect bullets, skip category labels ──
        if current_section == "KEY SKILLS":
            if stripped.startswith("•"):
                skills_buffer.append(stripped[1:].strip())
            # category subheadings (non-bullet) are intentionally dropped —
            # all skills render flat in the two-column table
            i += 1
            continue

        # ── Bullets (all sections except KEY SKILLS) ───────
        if stripped.startswith("•"):
            para_bullet(doc, stripped[1:].strip())
            last_was_date = False
            i += 1
            continue

        # ── PROFESSIONAL EXPERIENCE ────────────────────────
        if current_section == "PROFESSIONAL EXPERIENCE":
            # Location | Dates line
            if "|" in stripped:
                para_date_location(doc, stripped)
                last_was_date = True
                i += 1
                continue
            # Line with EM_DASH: role title (bold) OR thesis subtitle (italic)
            if EM_DASH in stripped:
                if last_was_date:
                    # Subtitle immediately after date line
                    para_subtitle(doc, stripped)
                else:
                    para_job_title(doc, stripped)
                last_was_date = False
                i += 1
                continue
            last_was_date = False

        # ── EDUCATION ─────────────────────────────────────
        if current_section == "EDUCATION":
            if "|" in stripped:
                p = doc.add_paragraph()
                add_run(p, stripped, size=10, italic=True, color=BLACK)
                sp(p, 0, 2)
                i += 1
                continue
            # Degree title: next non-empty line contains |
            next_non_empty = next(
                (lines[j].strip() for j in range(i + 1, len(lines)) if lines[j].strip()), ""
            )
            if "|" in next_non_empty:
                p = doc.add_paragraph()
                add_run(p, stripped, size=11, bold=True, color=BLACK)
                sp(p, 6, 0)
                i += 1
                continue
            # Thesis note
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_run(p, stripped, size=10, italic=True, color=BLACK)
            sp(p, 0, 4)
            i += 1
            continue

        # ── TECHNICAL SKILLS ──────────────────────────────
        if current_section == "TECHNICAL SKILLS":
            if ":" in stripped:
                colon = stripped.index(":")
                p     = doc.add_paragraph()
                add_run(p, stripped[: colon + 1], size=10.5, bold=True, color=BLACK)
                add_run(p, stripped[colon + 1 :], size=10.5, color=BLACK)
                sp(p, 0, 2)
                i += 1
                continue

        # ── ADDITIONAL EXPERIENCE ─────────────────────────
        if current_section == "ADDITIONAL EXPERIENCE":
            if EM_DASH in stripped and "|" not in stripped:
                para_job_title(doc, stripped)
                i += 1
                continue
            if not stripped.startswith("•"):
                p = doc.add_paragraph()
                add_run(p, stripped, size=10, italic=True, color=BLACK)
                sp(p, 0, 2)
                i += 1
                continue

        # ── LANGUAGES ─────────────────────────────────────
        if current_section == "LANGUAGES":
            p = doc.add_paragraph()
            add_run(p, stripped, size=10.5, color=BLACK)
            sp(p, 0, 2)
            i += 1
            continue

        # ── Default paragraph ──────────────────────────────
        p = doc.add_paragraph()
        add_run(p, stripped, size=10.5, color=BLACK)
        sp(p, 0, 3)
        i += 1

    # Flush any remaining skills
    if current_section == "KEY SKILLS":
        flush_skills()

    doc.save(str(docx_path))
    print(f"  Resume DOCX saved:       {docx_path.name}")


# ── Cover letter generator ────────────────────────────────

def generate_cover_letter_docx(md_path, docx_path):
    raw = Path(md_path).read_text(encoding="utf-8")

    # Extract English version only
    block = raw
    if "[English Version]" in raw:
        start = raw.index("[English Version]") + len("[English Version]")
        block = raw[start:]
    # Cut at German version or standalone "---" separator
    for marker in ("[Deutsche Version]", "\n---\n"):
        if marker in block:
            block = block[: block.index(marker)]

    lines = block.splitlines()

    doc = Document()
    set_margins(doc, 1.25)
    set_default_font(doc, size=11)

    for line in lines:
        stripped = line.strip()

        # Empty line → small spacer
        if not stripped:
            p = doc.add_paragraph()
            sp(p, 0, 4)
            continue

        # Skip the COVER LETTER title line
        if stripped.startswith("COVER LETTER"):
            continue

        # Subject line (contains "Application for" or "Bewerbung als")
        if stripped.startswith("Application for") or stripped.startswith("Bewerbung"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, stripped, size=11, bold=True, underline=True, color=BLACK)
            sp(p, 12, 12)
            continue

        # Bullet point
        if stripped.startswith("•"):
            para_bullet_with_bold_label(doc, stripped[1:].strip())
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_run(p, stripped, size=11, color=BLACK)
        sp(p, 0, 6)

    doc.save(str(docx_path))
    print(f"  Cover letter DOCX saved: {docx_path.name}")


# ── Entry point ───────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: /c/Python313/python.exe tools/make_docx.py <output_folder_path>")
        sys.exit(1)

    folder = Path(sys.argv[1])
    if not folder.exists():
        print(f"Folder not found: {folder}")
        sys.exit(1)

    print(f"Generating DOCX files in: {folder}")

    resume_md   = folder / "02-resume.md"
    resume_docx = folder / "02-resume.docx"
    cover_md    = folder / "03-cover-letter.md"
    cover_docx  = folder / "03-cover-letter.docx"

    if resume_md.exists():
        generate_resume_docx(resume_md, resume_docx)
    else:
        print(f"  Skipped resume (not found): {resume_md.name}")

    if cover_md.exists():
        generate_cover_letter_docx(cover_md, cover_docx)
    else:
        print(f"  Skipped cover letter (not found): {cover_md.name}")

    print("Done.")
